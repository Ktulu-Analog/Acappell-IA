"""
Utilitaires pour la génération de documents Word avec modèle personnalisé
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
import re


def load_template():
    """
    Charge le modèle Word si disponible, sinon crée un document vierge

    Returns:
        Document Word (avec ou sans modèle)
    """
    template_path = "modelerapport.docx"

    if os.path.exists(template_path):
        try:
            return Document(template_path)
        except Exception as e:
            print(f"Erreur lors du chargement du modèle : {e}")
            return Document()
    else:
        return Document()


def apply_custom_styles(doc):
    """
    Applique les styles personnalisés au document si le modèle n'est pas chargé
    """
    try:
        # Style pour le titre principal
        if 'CustomTitle' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(26)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 0, 0)
            title_style.paragraph_format.space_after = Pt(10)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Style pour les sous-titres
        if 'CustomSubtitle' not in [s.name for s in doc.styles]:
            subtitle_style = doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.name = 'Calibri'
            subtitle_style.font.size = Pt(14)
            subtitle_style.font.color.rgb = RGBColor(89, 89, 89)
            subtitle_style.paragraph_format.space_after = Pt(8)
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass  # Les styles existent déjà ou erreur


def add_logo_if_exists(doc):
    """Ajoute le logo si le fichier existe"""
    if os.path.exists("logo.png"):
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture("logo.png", width=Inches(2.36))
            paragraph.paragraph_format.space_after = Pt(12)
        except Exception as e:
            print(f"Erreur lors de l'ajout du logo : {e}")


def build_transcription_doc(segments, filename):
    """
    Construit un document Word avec la transcription

    Args:
        segments: Liste des segments de transcription
        filename: Nom du fichier audio source

    Returns:
        Document Word
    """
    doc = load_template()

    # Supprimer le contenu du modèle (garder uniquement les styles)
    for element in doc.element.body:
        if element.tag.endswith('p') or element.tag.endswith('tbl'):
            element.getparent().remove(element)

    apply_custom_styles(doc)

    # Logo
    add_logo_if_exists(doc)

    # Titre avec style personnalisé
    try:
        title = doc.add_paragraph("Transcription audio", style='TITRE')
    except:
        try:
            title = doc.add_paragraph("Transcription audio", style='CustomTitle')
        except:
            title = doc.add_heading("Transcription audio", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Métadonnées avec style corps de texte
    try:
        doc.add_paragraph(f"Fichier source : {filename}", style='Corps de texte')
        doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}", style='Corps de texte')
        doc.add_paragraph(f"Nombre de segments : {len(segments)}", style='Corps de texte')
    except:
        doc.add_paragraph(f"Fichier source : {filename}")
        doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Nombre de segments : {len(segments)}")

    doc.add_paragraph()  # Espace

    try:
        doc.add_heading("Transcription", level=2, style='Titre')
    except:
        doc.add_heading("Transcription", level=2)

    # Contenu
    for segment in segments:
        speaker = segment.get("speaker", "Inconnu")
        text = segment.get("text", "")

        p = doc.add_paragraph()

        # Locuteur en gras
        speaker_run = p.add_run(f"{speaker} : ")
        speaker_run.bold = True
        speaker_run.font.color.rgb = RGBColor(0, 51, 102)
        speaker_run.font.name = 'Calibri'

        # Texte
        text_run = p.add_run(text)
        text_run.font.size = Pt(11)
        text_run.font.name = 'Calibri'

    return doc


def parse_markdown_table(lines, start_idx):
    """
    Parse un tableau Markdown et retourne les données + l'index de fin

    Args:
        lines: Liste de toutes les lignes
        start_idx: Index de la première ligne du tableau

    Returns:
        (headers, rows, end_idx) ou (None, None, start_idx) si pas de tableau
    """
    if start_idx >= len(lines):
        return None, None, start_idx

    # Vérifier si c'est bien un tableau (présence de |)
    first_line = lines[start_idx].strip()
    if "|" not in first_line:
        return None, None, start_idx

    # Parser les en-têtes
    headers = [h.strip() for h in first_line.split("|") if h.strip()]

    # Vérifier la ligne de séparation (---)
    if start_idx + 1 >= len(lines):
        return None, None, start_idx

    sep_line = lines[start_idx + 1].strip()
    if not all(c in "-|: " for c in sep_line):
        return None, None, start_idx

    # Parser les lignes de données
    rows = []
    current_idx = start_idx + 2

    while current_idx < len(lines):
        line = lines[current_idx].strip()

        # Fin du tableau si ligne vide ou sans |
        if not line or "|" not in line:
            break

        # Parser la ligne
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            rows.append(cells)

        current_idx += 1

    return headers, rows, current_idx


def clean_html_tags(text):
    """
    Nettoie les balises HTML d'un texte

    Args:
        text: Texte contenant potentiellement des balises HTML

    Returns:
        Texte nettoyé avec les balises converties
    """
    if not text:
        return text

    # Remplacer les balises <br>, <br/>, <br /> par des retours à la ligne
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)

    # Supprimer d'autres balises HTML courantes si nécessaire
    text = re.sub(r'<p>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '', text, flags=re.IGNORECASE)

    return text


def add_formatted_text(paragraph, text, font_name='Calibri'):
    """
    Ajoute du texte avec formatage Markdown (gras, italique) à un paragraphe
    Gère également les retours à la ligne (incluant ceux issus de balises <br>)

    Supporte :
    - ***texte*** : gras + italique
    - **texte** : gras
    - *texte* : italique
    - \n : retour à la ligne

    Args:
        paragraph: Paragraphe Word
        text: Texte avec formatage Markdown
        font_name: Nom de la police à utiliser
    """
    if not text:
        return

    # Nettoyer les balises HTML
    text = clean_html_tags(text)

    # Diviser le texte par retours à la ligne
    lines = text.split('\n')

    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            # Ajouter un retour à la ligne entre les lignes
            paragraph.add_run('\n')

        # Parser le texte en identifiant tous les segments avec leur style
        segments = []
        pos = 0

        while pos < len(line):
            # Chercher le prochain pattern (en ordre de priorité : *** puis ** puis *)
            next_match = None
            next_style = None

            # Chercher ***
            match_3 = re.search(r"\*\*\*(.+?)\*\*\*", line[pos:])
            # Chercher **
            match_2 = re.search(r"\*\*(.+?)\*\*", line[pos:])
            # Chercher *
            match_1 = re.search(r"\*(.+?)\*", line[pos:])

            # Déterminer quel match utiliser (le plus proche, et en cas d'égalité le plus long)
            candidates = []
            if match_3:
                candidates.append((match_3.start(), 3, match_3, "bold_italic"))
            if match_2:
                candidates.append((match_2.start(), 2, match_2, "bold"))
            if match_1:
                candidates.append((match_1.start(), 1, match_1, "italic"))

            if not candidates:
                # Plus de formatage, ajouter le reste
                if pos < len(line):
                    segments.append(("normal", line[pos:]))
                break

            # Trier : d'abord par position de début, puis par longueur (plus long en premier)
            candidates.sort(key=lambda x: (x[0], -x[1]))
            start_pos, length, match, style = candidates[0]

            # Ajouter le texte avant le match
            if start_pos > 0:
                segments.append(("normal", line[pos : pos + start_pos]))

            # Ajouter le texte formaté
            segments.append((style, match.group(1)))

            # Avancer la position
            pos = pos + match.end()

        # Ajouter tous les segments au paragraphe
        for style, content in segments:
            run = paragraph.add_run(content)
            run.font.name = font_name

            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "bold_italic":
                run.bold = True
                run.italic = True


def add_formatted_paragraph(doc, text, style=None):
    """
    Ajoute un paragraphe avec formatage Markdown

    Args:
        doc: Document Word
        text: Texte avec formatage Markdown
        style: Style de paragraphe (optionnel)

    Returns:
        Paragraphe créé
    """
    p = doc.add_paragraph(style=style)
    add_formatted_text(p, text)
    return p


def build_summary_doc(summary_text, filename):
    """
    Construit un document Word avec la synthèse

    Args:
        summary_text: Texte de la synthèse généré par l'IA
        filename: Nom du fichier audio source

    Returns:
        Document Word
    """
    doc = load_template()

    # Supprimer le contenu du modèle (garder uniquement les styles)
    for element in doc.element.body:
        if element.tag.endswith('p') or element.tag.endswith('tbl'):
            element.getparent().remove(element)

    apply_custom_styles(doc)

    # Logo
    add_logo_if_exists(doc)

    # Titre avec style personnalisé
    try:
        title = doc.add_paragraph("Compte rendu de réunion", style='TITRE')
    except:
        try:
            title = doc.add_paragraph("Compte rendu de réunion", style='CustomTitle')
        except:
            title = doc.add_heading("Compte rendu de réunion", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Métadonnées
    try:
        doc.add_paragraph(f"Fichier source : {filename}", style='Sous-titre')
        doc.add_paragraph(
            f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            style='Sous-titre'
        )
    except:
        try:
            doc.add_paragraph(f"Fichier source : {filename}", style='CustomSubtitle')
            doc.add_paragraph(
                f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M')}",
                style='CustomSubtitle'
            )
        except:
            p1 = doc.add_paragraph(f"Fichier source : {filename}")
            p1.runs[0].font.size = Pt(12)
            p1.runs[0].font.color.rgb = RGBColor(89, 89, 89)
            p2 = doc.add_paragraph(
                f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            )
            p2.runs[0].font.size = Pt(12)
            p2.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()  # Espace

    # IMPORTANT : Vérifier que le texte n'est pas vide
    if not summary_text or not summary_text.strip():
        doc.add_paragraph("⚠️ Aucun contenu de synthèse généré.")
        return doc

    # Ajouter le contenu de la synthèse
    # Parser le texte ligne par ligne pour détecter les titres et le contenu
    lines = summary_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Ignorer les lignes vides
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Détecter les tableaux Markdown
        if "|" in line:
            headers, rows, end_idx = parse_markdown_table(lines, i)

            if headers and rows:
                # Créer un tableau Word
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"

                # En-têtes
                header_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    if idx < len(header_cells):
                        # Vider la cellule et ajouter le texte avec formatage
                        header_cells[idx].text = ""
                        p = header_cells[idx].paragraphs[0]
                        add_formatted_text(p, header)
                        # Mettre en gras les en-têtes
                        for run in p.runs:
                            run.bold = True
                            run.font.name = 'Calibri'

                # Données
                for row_idx, row_data in enumerate(rows):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(row_cells):
                            # Vider la cellule et ajouter le texte avec formatage
                            row_cells[col_idx].text = ""
                            p = row_cells[col_idx].paragraphs[0]
                            add_formatted_text(p, cell_data)

                doc.add_paragraph()  # Espace après le tableau
                i = end_idx
                continue

        # Détecter les titres (lignes commençant par #, ou en MAJUSCULES, ou suivies de :)
        if line.startswith("#"):
            # Format Markdown
            level = line.count("#", 0, 3)
            title_text = line.lstrip("#").strip()

            try:
                if level == 1:
                    doc.add_paragraph(title_text, style='Titre')
                else:
                    h = doc.add_heading(title_text, level=min(level, 3))
                    for run in h.runs:
                        run.font.name = 'Calibri'
            except:
                h = doc.add_heading(title_text, level=min(level, 3))
                for run in h.runs:
                    run.font.name = 'Calibri'

        elif line.isupper() and len(line) > 3 and len(line) < 100:
            # Ligne en majuscules = titre
            try:
                doc.add_paragraph(line.title(), style='Titre')
            except:
                h = doc.add_heading(line.title(), level=2)
                for run in h.runs:
                    run.font.name = 'Calibri'

        elif (
            line.endswith(":")
            and len(line) < 100
            and not line.startswith("-")
            and not line.startswith("*")
        ):
            # Ligne se terminant par : = sous-titre
            try:
                doc.add_paragraph(line.rstrip(":"), style='Sous-titre')
            except:
                h = doc.add_heading(line.rstrip(":"), level=3)
                for run in h.runs:
                    run.font.name = 'Calibri'

        elif line.startswith("- ") or line.startswith("* "):
            # Liste à puces
            text_content = line[2:]
            try:
                p = doc.add_paragraph(style="Puce")
            except:
                p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text_content)
            p.paragraph_format.left_indent = Inches(0.25)

        elif len(line) > 0 and line[0].isdigit() and len(line) > 2 and line[1] in ".):":
            # Liste numérotée
            text_content = line[2:].strip()
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, text_content)
            p.paragraph_format.left_indent = Inches(0.25)

        else:
            # Paragraphe normal
            try:
                p = doc.add_paragraph(style='Corps de texte')
            except:
                p = doc.add_paragraph()
            add_formatted_text(p, line)
            p.paragraph_format.line_spacing = 1.15

        i += 1

    return doc


def build_summary_doc_simple(summary_text, filename):
    """
    Version simple : ajoute tout le texte tel quel
    Utilisez cette fonction si la version complexe pose problème

    Args:
        summary_text: Texte de la synthèse
        filename: Nom du fichier source

    Returns:
        Document Word
    """
    doc = load_template()

    # Supprimer le contenu du modèle
    for element in doc.element.body:
        if element.tag.endswith('p') or element.tag.endswith('tbl'):
            element.getparent().remove(element)

    apply_custom_styles(doc)

    # Logo
    add_logo_if_exists(doc)

    # Titre
    try:
        title = doc.add_paragraph("Compte rendu de réunion", style='TITRE')
    except:
        try:
            title = doc.add_paragraph("Compte rendu de réunion", style='CustomTitle')
        except:
            title = doc.add_heading("Compte rendu de réunion", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Métadonnées
    try:
        doc.add_paragraph(f"Fichier source : {filename}", style='Sous-titre')
        doc.add_paragraph(
            f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            style='Sous-titre'
        )
    except:
        doc.add_paragraph(f"Fichier source : {filename}")
        doc.add_paragraph(
            f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )

    doc.add_paragraph()

    # Vérification
    if not summary_text or not summary_text.strip():
        doc.add_paragraph("⚠️ Aucun contenu de synthèse généré.")
        return doc

    # Nettoyer les balises HTML
    summary_text = clean_html_tags(summary_text)

    # Ajouter tout le texte
    for paragraph_text in summary_text.split("\n\n"):
        if paragraph_text.strip():
            try:
                p = doc.add_paragraph(paragraph_text.strip(), style='Corps de texte')
            except:
                p = doc.add_paragraph(paragraph_text.strip())
                p.runs[0].font.name = 'Calibri'

    return doc
